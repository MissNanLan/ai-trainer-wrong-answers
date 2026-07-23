"""Build the 1–300 AI trainer single-choice question bank from Word sources."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


SOURCE_ROOT = Path(__file__).resolve().parents[3]
QUESTION_FILE = SOURCE_ROOT / "人工智能训练师三级_单选题题目 1-300.docx"
ANSWER_FILES = [
    SOURCE_ROOT / "人工智能训练师三级_单选题1-100_答案解析.docx",
    SOURCE_ROOT / "人工智能训练师三级_单选题101-200_答案解析.docx",
    SOURCE_ROOT / "人工智能训练师三级_单选题201-300_答案解析.docx",
]

QUESTION_RE = re.compile(r"^(\d{1,3})[.．、]\s*(.+)$")
OPTION_RE = re.compile(r"^[（(]([ABCD])[）)]\s*(.+)$")
QUICK_ANSWER_RE = re.compile(r"(?<!\d)(\d{1,3})\s*[.．、]\s*([ABCD])(?=\s|$)")
DETAILED_ANSWER_RE = re.compile(r"^(\d{1,3})[.．、]\s*答案[:：]\s*([ABCD])$")


def nonempty_paragraphs(document: Document) -> list[str]:
    return [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]


def parse_questions() -> tuple[dict[int, dict], list[str]]:
    questions: dict[int, dict] = {}
    issues: list[str] = []
    current: dict | None = None

    for line in nonempty_paragraphs(Document(QUESTION_FILE)):
        question_match = QUESTION_RE.match(line)
        option_match = OPTION_RE.match(line)
        if question_match:
            if current:
                questions[current["id"]] = current
            current = {
                "id": int(question_match.group(1)),
                "topic": "人工智能训练师三级",
                "stem": question_match.group(2),
                "options": [],
            }
        elif option_match and current:
            current["options"].append({"key": option_match.group(1), "text": option_match.group(2)})
        elif current:
            current["stem"] += line

    if current:
        questions[current["id"]] = current

    for question_id, question in questions.items():
        if [option["key"] for option in question["options"]] != list("ABCD"):
            issues.append(f"题号 {question_id} 的选项不是完整 A–D")
    return questions, issues


def parse_answers() -> tuple[dict[int, dict], list[str]]:
    answers: dict[int, dict] = {}
    issues: list[str] = []
    for path in ANSWER_FILES:
        document = Document(path)
        for table in document.tables:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) < 3 or not cells[0].isdigit() or cells[1] not in "ABCD":
                    issues.append(f"{path.name} 存在无法识别的表格答案行：{' | '.join(cells)}")
                    continue
                answers[int(cells[0])] = {"answer": cells[1], "note": cells[2]}

        paragraphs = nonempty_paragraphs(document)
        for line in paragraphs:
            for number, answer in QUICK_ANSWER_RE.findall(line):
                answers[int(number)] = {"answer": answer, "note": ""}

        for index, line in enumerate(paragraphs):
            detailed = DETAILED_ANSWER_RE.match(line)
            if not detailed:
                continue
            question_id, answer = int(detailed.group(1)), detailed.group(2)
            note = paragraphs[index + 1].removeprefix("解析：") if index + 1 < len(paragraphs) else ""
            answers[question_id] = {"answer": answer, "note": note}

    return answers, issues


def parse_question_bank() -> tuple[list[dict], list[str]]:
    questions, issues = parse_questions()
    answers, answer_issues = parse_answers()
    issues.extend(answer_issues)
    output: list[dict] = []
    for question_id in range(1, 301):
        question = questions.get(question_id)
        answer = answers.get(question_id)
        if not question:
            issues.append(f"缺少题号 {question_id} 的题目")
            continue
        if not answer or answer["answer"] not in "ABCD":
            issues.append(f"缺少题号 {question_id} 的标准答案")
            continue
        output.append({
            **question,
            "answer": answer["answer"],
            "note": answer["note"] or "请结合题库答案解析复习。",
            "memory": answer["note"] or "请结合题库答案解析复习。",
        })
    return output, issues


def main() -> None:
    questions, issues = parse_question_bank()
    if issues:
        raise SystemExit("\n".join(issues))
    Path("question-bank-300.json").write_text(
        json.dumps(questions, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"已生成 {len(questions)} 道题：question-bank-300.json")


if __name__ == "__main__":
    main()
