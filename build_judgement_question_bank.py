"""Build the 1–300 AI trainer judgement-question bank from Word sources."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


QUESTION_NAME = "人工智能训练师三级_判断题题目 1-300.docx"
ANSWER_NAMES = [
    "人工智能训练师三级_判断题1-100_答案解析.docx",
    "人工智能训练师三级_判断题101-200_答案解析.docx",
    "人工智能训练师三级_判断题201-300_答案解析.docx",
]
QUESTION_RE = re.compile(r"^（\s*）\s*(\d{1,3})\.\s*(.+)$")
PARAGRAPH_ANSWER_RE = re.compile(r"^(\d{1,3})\.\s*答案：([√×])\s*解析：(.*)$", re.S)


def source_root() -> Path:
    for parent in Path(__file__).resolve().parents:
        if (parent / QUESTION_NAME).exists():
            return parent
    raise FileNotFoundError(f"找不到 {QUESTION_NAME}")


def parse_question_bank() -> tuple[list[dict], list[str]]:
    root = source_root()
    questions: dict[int, dict] = {}
    answers: dict[int, tuple[str, str]] = {}
    issues: list[str] = []

    for paragraph in Document(root / QUESTION_NAME).paragraphs:
        match = QUESTION_RE.match(paragraph.text.strip())
        if match:
            number = int(match.group(1))
            questions[number] = {"id": number, "topic": "人工智能训练师三级 判断题", "stem": match.group(2)}

    for name in ANSWER_NAMES[:2]:
        for row in Document(root / name).tables[0].rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 3 and cells[0].isdigit() and cells[1] in {"√", "×"}:
                answers[int(cells[0])] = (cells[1], cells[2])
            else:
                issues.append(f"{name} 存在无法识别的答案行")

    for paragraph in Document(root / ANSWER_NAMES[2]).paragraphs:
        match = PARAGRAPH_ANSWER_RE.match(paragraph.text.strip())
        if match:
            answers[int(match.group(1))] = (match.group(2), match.group(3).strip())

    output: list[dict] = []
    for number in range(1, 301):
        question, answer = questions.get(number), answers.get(number)
        if not question:
            issues.append(f"缺少题号 {number} 的题目")
            continue
        if not answer or answer[0] not in {"√", "×"} or not answer[1]:
            issues.append(f"缺少题号 {number} 的答案或解析")
            continue
        output.append({**question, "answer": answer[0], "note": answer[1], "memory": answer[1]})
    return output, issues


def main() -> None:
    questions, issues = parse_question_bank()
    if issues:
        raise SystemExit("\n".join(issues))
    Path("judgement-bank-300.json").write_text(json.dumps(questions, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"已生成 {len(questions)} 道判断题：judgement-bank-300.json")


if __name__ == "__main__":
    main()
